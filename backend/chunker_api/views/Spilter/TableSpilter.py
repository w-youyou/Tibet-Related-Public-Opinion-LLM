import os
import json
from dataclasses import dataclass
from typing import List, Dict, Any, Tuple

# 可选依赖：pdfplumber、python-docx、pandas、openpyxl
try:
    import pdfplumber  # 更稳健的 PDF 表格抽取
except Exception:
    pdfplumber = None

try:
    from docx import Document  # Word 读取
except Exception:
    Document = None

try:
    import pandas as pd  # Excel 读取
except Exception:
    pd = None


@dataclass
class TableChunk:
    content: List[List[str]]  # 包含表头行+数据行
    metadata: Dict[str, Any]


class TableSpilter:
    """表格分块器：识别表头后，每30行作为一个块，每块都带表头。
    支持 PDF（pdfplumber）、Word（python-docx）、Excel（pandas/openpyxl）。
    元数据包含：file_name, table_name。
    """

    def __init__(self, rows_per_chunk: int = 30):
        self.rows_per_chunk = rows_per_chunk

    # ---------------- 核心分块逻辑 ---------------- #
    def split_table_rows(self, rows: List[List[Any]], file_name: str, table_name: str,
                         override_header: List[str] = None, override_table_name: str = None, 
                         file_path: str = None) -> List[TableChunk]:
        if not rows:
            return []
        # 预处理：统一列数、清洗文本并进行扁平化（行/列方向前向填充以应对合并单元格）
        rows = self._normalize_and_flatten(rows)
        # 识别表头：第一行（可被覆盖）
        header = ["" if v is None else str(v).strip() for v in rows[0]]
        header = self._make_header_unique(header)
        if override_header is not None and any(override_header):
            header = ["" if v is None else str(v).strip() for v in override_header]
            header = self._make_header_unique(header)
        data_rows = [
            ["" if v is None else str(v).strip() for v in r]
            for r in rows[1:]
        ]

        chunks: List[TableChunk] = []
        fixed_table_name = override_table_name if (override_table_name and str(override_table_name).strip()) else table_name
        for i in range(0, len(data_rows), self.rows_per_chunk):
            part = data_rows[i:i + self.rows_per_chunk]
            content = [header] + part
            # 构建 metadata，过滤掉 None 值（Chroma 不接受 None）
            metadata = {}
            if file_name is not None:
                metadata["file_name"] = file_name
            if fixed_table_name is not None:
                metadata["table_name"] = fixed_table_name
            # 如果提供了文件路径，添加到元数据中
            if file_path is not None:
                metadata["source_file"] = file_path
            chunks.append(TableChunk(
                content=content,
                metadata=metadata
            ))
        return chunks

    def _normalize_and_flatten(self, rows: List[List[Any]]) -> List[List[str]]:
        # 统一列数
        col_count = max(len(r) for r in rows)
        norm = [
            ["" if v is None else str(v).strip().replace("\u3000", " ").replace("\n", " ") for v in (r + [None] * (col_count - len(r)))]
            for r in rows
        ]
        norm = [r for r in norm]
        # 行内前向填充（处理横向合并）
        for r in norm:
            last = ""
            for j in range(col_count):
                if r[j] == "":
                    r[j] = last
                else:
                    last = r[j]
        # 列向前向填充（处理纵向合并），从第二行开始保留表头
        for j in range(col_count):
            last = norm[0][j]
            for i in range(1, len(norm)):
                if norm[i][j] == "":
                    norm[i][j] = last
                else:
                    last = norm[i][j]
        return norm

    def _make_header_unique(self, header: List[str]) -> List[str]:
        seen: Dict[str, int] = {}
        result: List[str] = []
        for h in header:
            key = h or "col"
            cnt = seen.get(key, 0)
            if cnt == 0:
                result.append(key)
            else:
                result.append(f"{key}_{cnt}")
            seen[key] = cnt + 1
        return result

    # ---------------- JSON 记录转换 ---------------- #
    def rows_to_records(self, rows: List[List[str]]) -> List[Dict[str, Any]]:
        if not rows:
            return []
        col_count = max(len(r) for r in rows)
        header = rows[0] + [""] * (col_count - len(rows[0]))
        header = self._make_header_unique(header)
        records: List[Dict[str, Any]] = []
        for r in rows[1:]:
            row = r + [""] * (col_count - len(r))
            obj = {header[i]: row[i] for i in range(col_count)}
            records.append(obj)
        return records

    # ---------------- 数据源读取 ---------------- #
    def read_pdf_tables(self, file_path: str) -> List[Tuple[str, List[List[Any]]]]:
        """返回 [(table_name, rows)] 列表；table_name 为表格上一段文字。
        修复：记录第一个表格的表头，后续页面复用该表头。
        """
        if pdfplumber is None:
            raise ImportError("缺少依赖：pdfplumber。请安装 pip install pdfplumber")
        tables: List[Tuple[str, List[List[Any]]]] = []
        file_name = os.path.basename(file_path)
        
        # 记录第一个表格的表头，用于后续页面复用
        global_header = None
        global_table_name = None
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                # 提取页面表格
                page_tables = page.extract_tables() or []
                # 近似作为表名：取表格上方最近的一行文本
                page_text = page.extract_text() or ""
                text_lines = [l.strip() for l in page_text.splitlines() if l.strip()]
                # 采用简单启发：表名设为该页首个非空文本，若无则文件名
                default_table_name = text_lines[0] if text_lines else file_name
                
                for tbl in page_tables:
                    # 过滤空行
                    rows = [list(map(lambda x: "" if x is None else str(x), r)) for r in tbl if any(c is not None and str(c).strip() for c in r)]
                    if len(rows) >= 1 and any(cell.strip() for cell in rows[0]):
                        # 如果是第一个表格，记录表头
                        if global_header is None:
                            global_header = rows[0].copy()  # 深拷贝表头
                            global_table_name = default_table_name
                            tables.append((default_table_name, rows))
                        else:
                            # 后续表格：使用全局表头，只保留数据行
                            # 检查列数是否匹配
                            if len(rows[0]) == len(global_header):
                                # 使用全局表头 + 当前页面的数据行
                                merged_rows = [global_header] + rows[1:]  # 跳过当前页面的表头
                                tables.append((global_table_name, merged_rows))
                            else:
                                # 列数不匹配，作为新表格处理
                                global_header = rows[0].copy()
                                global_table_name = default_table_name
                                tables.append((default_table_name, rows))
        return tables

    def read_docx_tables(self, file_path: str) -> List[Tuple[str, List[List[Any]]]]:
        if Document is None:
            raise ImportError("缺少依赖：python-docx。请安装 pip install python-docx")
        doc = Document(file_path)
        tables: List[Tuple[str, List[List[Any]]]] = []
        file_name = os.path.basename(file_path)
        # 简单策略：表格前最后一段非空文字作为表名
        paragraphs_text = [p.text.strip() for p in doc.paragraphs]
        non_empty_idx = [i for i, t in enumerate(paragraphs_text) if t]
        cur_para_idx = 0
        for t in doc.tables:
            rows: List[List[str]] = []
            for row in t.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            # 找到最近的上一非空段落
            table_name = file_name
            while cur_para_idx < len(non_empty_idx) and non_empty_idx[cur_para_idx] < len(paragraphs_text):
                cur_para_idx += 1
            for i in range(len(paragraphs_text) - 1, -1, -1):
                if paragraphs_text[i]:
                    table_name = paragraphs_text[i]
                    break
            if rows:
                tables.append((table_name, rows))
        return tables

    def read_excel_tables(self, file_path: str) -> List[Tuple[str, List[List[Any]]]]:
        if pd is None:
            raise ImportError("缺少依赖：pandas/openpyxl。请安装 pip install pandas openpyxl")
        tables: List[Tuple[str, List[List[Any]]]] = []
        file_name = os.path.basename(file_path)
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            # 使用工作表第一行作为表头：不让 pandas 吃掉首行
            df = xls.parse(sheet_name, header=None)
            rows = [["" if pd.isna(v) else str(v) for v in row] for row in df.values.tolist()]
            # 识别真正的表头行：在前20行内选择非空单元格最多且不全相同的一行
            header_idx = 0
            max_non_empty = -1
            for i in range(min(20, len(rows))):
                row = rows[i]
                non_empty_vals = [c for c in row if str(c).strip() != ""]
                if not non_empty_vals:
                    continue
                # 跳过整行相同（例如标题跨列合并被填充）的情况
                if len(set(non_empty_vals)) == 1:
                    continue
                if len(non_empty_vals) > max_non_empty:
                    max_non_empty = len(non_empty_vals)
                    header_idx = i
            # 将识别的表头置顶，并截取该行及其后的数据行为表
            if rows:
                rows = rows[header_idx:]
            if rows:
                tables.append((sheet_name or file_name, rows))
        return tables

    def read_doc_tables(self, file_path: str) -> List[Tuple[str, List[List[Any]]]]:
        """读取.doc文件中的表格"""
        try:
            import win32com.client
            import pythoncom
            import os
            
            # 转换为绝对路径
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                print(f"文件不存在: {abs_path}")
                return []
            
            # 初始化COM组件
            pythoncom.CoInitialize()
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(abs_path)
            tables: List[Tuple[str, List[List[Any]]]] = []
            file_name = os.path.basename(file_path)
            
            # 遍历文档中的所有表格
            for i, table in enumerate(doc.Tables):
                rows: List[List[str]] = []
                for row in table.Rows:
                    row_data = []
                    for cell in row.Cells:
                        cell_text = cell.Range.Text.strip()
                        # 移除Word表格中的特殊字符
                        cell_text = cell_text.replace('\r', '').replace('\a', '')
                        row_data.append(cell_text)
                    if any(cell.strip() for cell in row_data):  # 只添加非空行
                        rows.append(row_data)
                
                if rows:
                    table_name = f"表格_{i+1}"
                    tables.append((table_name, rows))
            
            doc.Close()
            word.Quit()
            
            return tables
        except Exception as e:
            print(f"读取.doc表格失败: {e}")
            return []
        finally:
            try:
                if 'word' in locals():
                    word.Quit()
                # 清理COM组件
                pythoncom.CoUninitialize()
            except:
                pass

    def read_any(self, file_path: str) -> List[Tuple[str, List[List[Any]]]]:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.read_pdf_tables(file_path)
        if ext == ".docx":
            return self.read_docx_tables(file_path)
        if ext == ".doc":
            return self.read_doc_tables(file_path)
        if ext in {".xls", ".xlsx"}:
            return self.read_excel_tables(file_path)
        raise ValueError(f"不支持的文件类型: {ext}")

    # ---------------- 输出与测试 ---------------- #
    def save_to_jsonl(self, chunks: List[TableChunk], output_path: str, source_file_path: str = None) -> None:
        """将表格分块结果保存为 JSONL（每行一个 JSON）。
        
        Args:
            chunks: 表格分块列表
            output_path: 输出文件路径
            source_file_path: 源文件路径（可选，用于元数据）
        """
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            for i, c in enumerate(chunks, 1):
                content_json = self.rows_to_records(c.content)
                metadata = c.metadata.copy()
                # 如果提供了源文件路径且元数据中没有，则添加
                if source_file_path and "source_file" not in metadata:
                    metadata["source_file"] = source_file_path
                obj = {
                    "id": i,
                    # 将内容序列化为字符串，但保持值内容不变
                    "content": json.dumps(content_json, ensure_ascii=False),
                    "metadata": metadata,
                }
                f.write(json.dumps(obj, ensure_ascii=False) + "\n")


def main():
    # 测试文件：PDF
    file_path = os.path.join(
        "知识库",
        "4 表格类文本",
        "6.涉藏舆情政策文件示例清单.pdf",
    )
    splitter = TableSpilter(rows_per_chunk=30)
    file_name = os.path.basename(file_path)

    try:
        tables = splitter.read_any(file_path)
    except Exception as e:
        print(f"读取表格失败: {e}")
        return

    all_chunks: List[TableChunk] = []
    # 按表（Excel 的各个 Sheet）分别确定并复用各自表头与表名
    for table_name, rows in tables:
        local_header: List[str] = []
        if rows and rows[0]:
            local_header = ["" if v is None else str(v).strip() for v in rows[0]]
        chunks = splitter.split_table_rows(
            rows,
            file_name=file_name,
            table_name=table_name,
            override_header=local_header,
            override_table_name=table_name,
            file_path=file_path,
        )
        all_chunks.extend(chunks)

    print(f"共生成分块: {len(all_chunks)}")
    # 简要预览前2个块（显示首条JSON记录）
    for i, c in enumerate(all_chunks[:2], 1):
        recs = splitter.rows_to_records(c.content)
        preview = json.dumps(recs[:1], ensure_ascii=False)
        print(f"块 {i} | 表名: {c.metadata.get('table_name')}\n示例: {preview}\n...")

    output_path = os.path.join("output", "表格分块.jsonl")
    print(f"保存到: {output_path}")
    splitter.save_to_jsonl(all_chunks, output_path, file_path)


if __name__ == "__main__":
    main()


