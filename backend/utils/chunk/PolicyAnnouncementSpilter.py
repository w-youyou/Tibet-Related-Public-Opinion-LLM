import re
import json
import os
from datetime import datetime
from typing import List, Dict, Any
from dataclasses import dataclass, asdict
from docx import Document
import fitz  # PyMuPDF for PDF processing
import pandas as pd  # for table extraction


@dataclass
class TextChunk:
    """文本分块结果"""
    content: str
    metadata: Dict[str, Any] = None  # 元数据信息
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class PolicyAnnouncementSpilter:
    """政策公告类文本分块器
    
    专门用于处理政策公告类文档，基于一级标题进行分块：
    - 一级标题：一、二、三、等
    """
    
    def __init__(self):
        # 定义分块模式，只保留一级标题
        self.patterns = [
            # 一级标题：一、二、三、等
            (r'^[一二三四五六七八九十百千万]+、', 1, '一级标题'),
            # 列表项：(一)、1. 等
            (r'^\s*\(?[一二三四五六七八九十百千万\d]+\)|[\d]+\.|[\-•] ', 2, '列表项'),
        ]
        
        # 编译正则表达式
        self.compiled_patterns = [
            (re.compile(pattern), level, name) 
            for pattern, level, name in self.patterns
        ]
    
    def split_text(self, text: str, tables_info: List[Dict[str, Any]] = None, file_path: str = None) -> List[TextChunk]:
        """将文本按政策公告格式进行分块
        
        Args:
            text: 待分块的文本
            tables_info: 表格信息列表
            
        Returns:
            分块结果列表
        """
        if not text or not text.strip():
            return []
        
        if tables_info is None:
            tables_info = []
        
        # 识别大标题、发布时间和介绍性内容
        main_title, publish_time, intro = self._extract_main_info(text)
        
        # 按行分割文本
        lines = text.split('\n')
        chunks = []
        current_chunk = []
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                if current_chunk:
                    current_chunk.append("")
                continue
            
            # 检查是否是表格标记
            if line.startswith("[TABLE_"):
                # 如果当前有未完成的分块，先保存它
                if current_chunk:
                    chunk_content = '\n'.join(current_chunk).strip()
                    if chunk_content:
                        # 构建 metadata，过滤掉 None 值（Chroma 不接受 None）
                        metadata = {}
                        if main_title is not None:
                            metadata["announcement_title"] = main_title
                        if publish_time is not None:
                            metadata["publish_time"] = publish_time
                        if intro is not None:
                            metadata["intro"] = intro
                        # 如果提供了文件路径，添加到元数据中
                        if file_path is not None:
                            metadata["source_file"] = file_path
                        chunks.append(TextChunk(
                            content=chunk_content,
                            metadata=metadata
                        ))
                
                # 处理表格
                table_index = int(line.replace("[TABLE_", "").replace("]", ""))
                if table_index < len(tables_info):
                    table_info = tables_info[table_index]
                    # 构建 metadata，过滤掉 None 值（Chroma 不接受 None）
                    metadata = {}
                    if main_title is not None:
                        metadata["announcement_title"] = main_title
                    if publish_time is not None:
                        metadata["publish_time"] = publish_time
                    if intro is not None:
                        metadata["intro"] = intro
                    if table_info.get('table_title') is not None:
                        metadata["table_title"] = table_info['table_title']
                    # 如果提供了文件路径，添加到元数据中
                    if file_path is not None:
                        metadata["source_file"] = file_path
                    chunks.append(TextChunk(
                        content=table_info['content'],
                        metadata=metadata
                    ))
                
                current_chunk = []
                continue
            
            # 检查是否匹配任何分块模式
            matched = False
            for pattern, level, pattern_name in self.compiled_patterns:
                match = pattern.match(line)
                if match:
                    # 如果当前有未完成的分块，先保存它
                    if current_chunk:
                        chunk_content = '\n'.join(current_chunk).strip()
                        if chunk_content:
                            # 构建 metadata，过滤掉 None 值（Chroma 不接受 None）
                            metadata = {}
                            if main_title is not None:
                                metadata["announcement_title"] = main_title
                            if publish_time is not None:
                                metadata["publish_time"] = publish_time
                            if intro is not None:
                                metadata["intro"] = intro
                            # 如果提供了文件路径，添加到元数据中
                            if file_path is not None:
                                metadata["source_file"] = file_path
                            chunks.append(TextChunk(
                                content=chunk_content,
                                metadata=metadata
                            ))
                    
                    # 开始新的分块
                    current_chunk = [line]
                    matched = True
                    break
            
            if not matched:
                # 如果当前行不匹配任何模式，添加到当前分块
                current_chunk.append(line)
        
        # 处理最后一个分块
        if current_chunk:
            chunk_content = '\n'.join(current_chunk).strip()
            if chunk_content:
                # 构建 metadata，过滤掉 None 值（Chroma 不接受 None）
                metadata = {}
                if main_title is not None:
                    metadata["announcement_title"] = main_title
                if publish_time is not None:
                    metadata["publish_time"] = publish_time
                if intro is not None:
                    metadata["intro"] = intro
                # 如果提供了文件路径，添加到元数据中
                if file_path is not None:
                    metadata["source_file"] = file_path
                chunks.append(TextChunk(
                    content=chunk_content,
                    metadata=metadata
                ))
        
        return chunks
    
    def _extract_main_info(self, text: str) -> tuple[str, str, str]:
        """提取大标题、发布时间和介绍性内容
        
        Args:
            text: 文本内容
            
        Returns:
            (大标题, 发布时间, 介绍性内容)
        """
        lines = text.split('\n')
        main_title = ""
        publish_time = ""
        intro_content = []
        # 提取大标题（第一行）
        for line in lines:
            line = line.strip()
            if line:
                main_title = line
                break
        
        # 提取发布时间和介绍性内容
        found_publish_time = False
        found_first_level = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 识别发布时间
            if "发布日期" in line or "发布时间" in line or "日期" in line:
                # 提取日期部分
                import re
                date_pattern = r'(\d{4}[-年]\d{1,2}[-月]\d{1,2}[日]?)'
                match = re.search(date_pattern, line)
                if match:
                    publish_time = match.group(1)
                    found_publish_time = True
                continue
            
            # 检查是否遇到第一个一级标题
            if not found_first_level:
                for pattern, level, pattern_name in self.compiled_patterns:
                    if level == 1 and pattern.match(line):
                        found_first_level = True
                        break
                
                # 如果已经找到发布时间但还没遇到一级标题，收集介绍性内容
                if found_publish_time and not found_first_level:
                    intro_content.append(line)
            else:
                # 已经遇到一级标题，停止收集介绍性内容
                break
        
        intro = '\n'.join(intro_content).strip() if intro_content else ""
        return main_title, publish_time, intro
    
    def get_chunks_by_level(self, chunks: List[TextChunk], level: int) -> List[TextChunk]:
        """获取指定层级的分块（已弃用，保留兼容性）
        
        Args:
            chunks: 分块列表
            level: 目标层级
            
        Returns:
            空列表（因为不再有层级概念）
        """
        return []
    
    def get_chunk_hierarchy(self, chunks: List[TextChunk]) -> Dict[str, Any]:
        """获取分块的结构信息
        
        Args:
            chunks: 分块列表
            
        Returns:
            结构信息字典
        """
        hierarchy = {
            "total_chunks": len(chunks),
            "chunks": []
        }
        
        for i, chunk in enumerate(chunks):
            hierarchy["chunks"].append({
                'id': i + 1,
                'content_preview': chunk.content[:100] + "..." if len(chunk.content) > 100 else chunk.content,
                'announcement_title': chunk.metadata.get('announcement_title', ''),
                'publish_time': chunk.metadata.get('publish_time', '')
            })
        
        return hierarchy
    
    def read_docx(self, file_path: str) -> tuple[str, List[Dict[str, Any]]]:
        """读取Word文档内容（包括段落和表格）
        
        Args:
            file_path: Word文档路径
            
        Returns:
            (文档文本内容, 表格信息列表)
        """
        try:
            doc = Document(file_path)
            text_content = []
            tables_info = []
            
            # 按文档顺序读取所有元素（段落和表格）
            for element in doc.element.body:
                if element.tag.endswith('p'):  # 段落
                    paragraph = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            paragraph = p
                            break
                    if paragraph and paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                elif element.tag.endswith('tbl'):  # 表格
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    if table:
                        # 处理表格内容
                        table_text = self._extract_table_text(table)
                        if table_text:
                            # 获取表格标题（上一行内容）
                            table_title = ""
                            if text_content:
                                table_title = text_content[-1]
                            
                            # 将表格信息存储到列表中
                            tables_info.append({
                                'content': table_text,
                                'table_title': table_title,
                                'position': len(text_content)  # 记录在文档中的位置
                            })
                            
                            # 在文本中插入表格标记
                            text_content.append(f"[TABLE_{len(tables_info)-1}]")
            
            return '\n'.join(text_content), tables_info
        except Exception as e:
            print(f"读取Word文档失败: {e}")
            return "", []
    
    def _extract_table_text(self, table) -> str:
        """提取表格文本内容
        
        Args:
            table: Word表格对象
            
        Returns:
            表格的文本表示
        """
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            
            if row_text:
                # 用制表符分隔单元格内容
                table_text.append('\t'.join(row_text))
        
        # 用换行符分隔各行
        return '\n'.join(table_text)
    
    def read_pdf(self, file_path: str) -> tuple[str, List[Dict[str, Any]]]:
        """读取PDF文档内容（包括文本和表格）
        
        Args:
            file_path: PDF文档路径
            
        Returns:
            (文档文本内容, 表格信息列表)
        """
        try:
            doc = fitz.open(file_path)
            text_content = []
            tables_info = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # 提取文本内容
                page_text = page.get_text()
                if page_text.strip():
                    text_content.append(page_text.strip())
                
                # 提取表格
                tables = page.find_tables()
                for table_idx, table in enumerate(tables):
                    try:
                        # 提取表格数据
                        table_data = table.extract()
                        if table_data:
                            # 将表格数据转换为文本格式
                            table_text = self._convert_table_data_to_text(table_data)
                            
                            # 获取表格标题（上一行内容）
                            table_title = ""
                            if text_content:
                                # 从当前页面的文本中查找表格前的标题
                                page_lines = page_text.split('\n')
                                for i, line in enumerate(page_lines):
                                    if table_text.split('\n')[0] in line:
                                        if i > 0:
                                            table_title = page_lines[i-1].strip()
                                        break
                            
                            # 将表格信息存储到列表中
                            tables_info.append({
                                'content': table_text,
                                'table_title': table_title,
                                'position': len(text_content) - 1  # 记录在文档中的位置
                            })
                            
                            # 在文本中插入表格标记
                            text_content.append(f"[TABLE_{len(tables_info)-1}]")
                    except Exception as e:
                        print(f"提取表格失败: {e}")
                        continue
            
            doc.close()
            return '\n'.join(text_content), tables_info
        except Exception as e:
            print(f"读取PDF文档失败: {e}")
            return "", []
    
    def _convert_table_data_to_text(self, table_data: List[List[str]]) -> str:
        """将表格数据转换为文本格式
        
        Args:
            table_data: 表格数据列表
            
        Returns:
            表格的文本表示
        """
        table_text = []
        
        for row in table_data:
            if row:  # 确保行不为空
                # 过滤掉空值并转换为字符串
                row_text = [str(cell).strip() if cell else "" for cell in row]
                # 用制表符分隔单元格内容
                table_text.append('\t'.join(row_text))
        
        # 用换行符分隔各行
        return '\n'.join(table_text)
    
    def read_doc(self, file_path: str) -> tuple[str, List[Dict[str, Any]]]:
        """读取.doc文档内容（包括文本和表格）
        
        Args:
            file_path: .doc文档路径
            
        Returns:
            (文档文本内容, 表格信息列表)
        """
        try:
            import win32com.client
            import pythoncom
            import os
            
            # 转换为绝对路径
            abs_path = os.path.abspath(file_path)
            if not os.path.exists(abs_path):
                print(f"文件不存在: {abs_path}")
                return "", []
            
            # 初始化COM组件
            pythoncom.CoInitialize()
            
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # 使用绝对路径打开文档
            doc = word.Documents.Open(abs_path)
            
            # 按段落读取，保持段落结构
            paragraphs = []
            for paragraph in doc.Paragraphs:
                text = paragraph.Range.Text.strip()
                if text:  # 只添加非空段落
                    paragraphs.append(text)
            
            doc.Close()
            word.Quit()
            
            # 用换行符连接段落
            text_content = '\n'.join(paragraphs)
            
            # .doc文件暂不支持表格提取，返回空列表
            tables_info = []
            return text_content, tables_info
        except Exception as e:
            print(f"读取.doc文档失败: {e}")
            return "", []
        finally:
            try:
                if 'word' in locals():
                    word.Quit()
                # 清理COM组件
                pythoncom.CoUninitialize()
            except:
                pass
    
    def read_document(self, file_path: str) -> tuple[str, List[Dict[str, Any]]]:
        """读取文档内容（支持Word和PDF）
        
        Args:
            file_path: 文档路径
            
        Returns:
            (文档文本内容, 表格信息列表)
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self.read_docx(file_path)
        elif file_ext == '.doc':
            return self.read_doc(file_path)
        elif file_ext == '.pdf':
            return self.read_pdf(file_path)
        else:
            print(f"不支持的文件格式: {file_ext}")
            return "", []
    
    def save_to_jsonl(self, chunks: List[TextChunk], output_path: str, document_metadata: Dict[str, Any] = None):
        """将分块结果保存为JSONL格式
        
        Args:
            chunks: 分块列表
            output_path: 输出文件路径
            document_metadata: 文档元数据
        """
        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for i, chunk in enumerate(chunks):
                # 构建完整的元数据
                chunk_data = {
                    "id": i + 1,
                    "content": chunk.content,
                    "metadata": {
                        **chunk.metadata,
                        "file_title": document_metadata.get("document_title", "") if document_metadata else ""
                    }
                }
                
                # 写入JSONL格式（每行一个JSON对象）
                f.write(json.dumps(chunk_data, ensure_ascii=False) + '\n')
        
        print(f"分块结果已保存到: {output_path}")
        print(f"共保存 {len(chunks)} 个分块")
    
    def print_chunks(self, chunks: List[TextChunk], show_content: bool = False):
        """打印分块结果
        
        Args:
            chunks: 分块列表
            show_content: 是否显示完整内容
        """
        print(f"共找到 {len(chunks)} 个分块：\n")
        
        for i, chunk in enumerate(chunks, 1):
            print(f"{i}. 分块 {i}")
            
            if show_content:
                content_lines = chunk.content.split('\n')
                for line in content_lines[:3]:  # 只显示前3行
                    print(f"   {line}")
                if len(content_lines) > 3:
                    print(f"   ...")
            print()


def main():
    print("开始测试")
    """测试函数"""
    
    # Word文档路径
    docx_path = r"知识库\政策文件\医保\兰医保〔2023〕68号-关于转发《关于印发甘肃省医疗保障政务服务事项清单2023版的通知》的通知---附件.docx"
    output_path = "./output/医保.jsonl"
    
    # 创建分块器
    splitter = PolicyAnnouncementSpilter()
    
    # 读取文档
    print(f"正在读取文档: {docx_path}")
    text_content, tables_info = splitter.read_document(docx_path)
    
    # 执行分块
    print("正在执行文本分块...")
    chunks = splitter.split_text(text_content, tables_info, docx_path)
    
    # 打印结果
    print("=== 分块结果 ===")
    splitter.print_chunks(chunks, show_content=True)
    
    # 分块结构
    print("\n=== 分块结构 ===")
    hierarchy = splitter.get_chunk_hierarchy(chunks)
    print(f"总分块数: {hierarchy['total_chunks']}")
    for chunk_info in hierarchy['chunks']:
        print(f"  - 分块 {chunk_info['id']}: {chunk_info['content_preview']}")
        print(f"    公告标题: {chunk_info['announcement_title']}")
        print(f"    发布时间: {chunk_info['publish_time']}")
        print()
    
    # 准备文档元数据
    # 从文件路径中提取文件名（不含扩展名）作为document_title
    import os
    file_name = os.path.basename(docx_path)
    document_title = os.path.splitext(file_name)[0]  # 去掉扩展名
    
    document_metadata = {
        "document_title": document_title
    }
    
    # 保存为JSONL格式
    print(f"\n正在保存结果到: {output_path}")
    splitter.save_to_jsonl(chunks, output_path, document_metadata)


if __name__ == "__main__":
    main()
