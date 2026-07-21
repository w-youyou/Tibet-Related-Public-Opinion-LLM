#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
基础文本分块器
提供多种基础分块策略：按长度切分、按标点符号切分、智能递归切分、按页切分
"""

import os
import re
from typing import List, Optional
from dataclasses import dataclass

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import pdfplumber
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_SUPPORT = True
except ImportError:
    PYMUPDF_SUPPORT = False


@dataclass
class Chunk:
    """分块数据类"""
    id: int
    content: str
    metadata: dict


class BasicSpilter:
    """基础文本分块器"""
    
    def __init__(self):
        """初始化基础分块器"""
        pass
    
    def read_text_file(self, file_path: str) -> str:
        """读取纯文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise ValueError(f"无法读取文件: {file_path}")
    
    def read_pdf(self, file_path: str) -> str:
        """读取PDF文件文本"""
        if not PDF_SUPPORT:
            raise ImportError("需要安装 pdfplumber 或 PyPDF2")
        
        text_content = []
        
        try:
            # 优先使用 pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
        except:
            # 备用方案：使用 PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
            except Exception as e:
                raise ValueError(f"读取PDF失败: {e}")
        
        return '\n'.join(text_content)
    
    def read_docx(self, file_path: str) -> str:
        """读取DOCX文件文本"""
        if not DOCX_SUPPORT:
            raise ImportError("需要安装 python-docx")
        
        try:
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return '\n'.join(paragraphs)
        except Exception as e:
            raise ValueError(f"读取DOCX失败: {e}")
    
    def get_file_text(self, file_path: str) -> str:
        """根据文件类型读取文本"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.txt' or ext == '.md':
            return self.read_text_file(file_path)
        elif ext == '.pdf':
            return self.read_pdf(file_path)
        elif ext == '.docx':
            return self.read_docx(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")
    
    # ==================== 按长度切分 ====================
    
    def split_by_length(self, 
                       text: str, 
                       chunk_size: int = 1000, 
                       chunk_overlap: int = 200) -> List[Chunk]:
        """
        按固定长度切分文本
        
        Args:
            text: 输入文本
            chunk_size: 每块字符数
            chunk_overlap: 重叠字符数
        
        Returns:
            分块列表
        """
        if not text.strip():
            return []
        
        chunks = []
        start = 0
        chunk_id = 1
        
        while start < len(text):
            end = min(start + chunk_size, len(text))
            chunk_text = text[start:end]
            
            if chunk_text.strip():
                chunks.append(Chunk(
                    id=chunk_id,
                    content=chunk_text.strip(),
                    metadata={
                        'chunk_method': 'by_length',
                        'chunk_size': chunk_size,
                        'chunk_overlap': chunk_overlap,
                        'start_pos': start,
                        'end_pos': end
                    }
                ))
                chunk_id += 1
            
            start = end - chunk_overlap if end < len(text) else len(text)
        
        return chunks
    
    # ==================== 按标点符号切分 ====================
    
    def split_by_punctuation(self, 
                             text: str, 
                             chunk_size: int = 1000,
                             chunk_overlap: int = 200) -> List[Chunk]:
        """
        按标点符号切分文本
        
        Args:
            text: 输入文本
            chunk_size: 目标每块字符数
            chunk_overlap: 重叠字符数
        
        Returns:
            分块列表
        """
        if not text.strip():
            return []
        
        # 定义标点符号模式（中文和英文）
        punctuation_pattern = r'[。！？\n\n]+|[.!?\n\n]+'
        
        # 按标点符号分割文本
        segments = re.split(punctuation_pattern, text)
        segments = [s.strip() for s in segments if s.strip()]
        
        if not segments:
            return []
        
        chunks = []
        current_chunk = []
        current_length = 0
        chunk_id = 1
        start_pos = 0
        
        for i, segment in enumerate(segments):
            segment_length = len(segment)
            
            # 如果当前块加上新段落后超过目标大小，且当前块不为空
            if current_length + segment_length > chunk_size and current_chunk:
                # 创建当前块
                chunk_text = ' '.join(current_chunk)
                chunks.append(Chunk(
                    id=chunk_id,
                    content=chunk_text,
                    metadata={
                        'chunk_method': 'by_punctuation',
                        'chunk_size': chunk_size,
                        'chunk_overlap': chunk_overlap,
                        'segment_count': len(current_chunk)
                    }
                ))
                chunk_id += 1
                
                # 保留重叠部分（从后往前取几个段落）
                overlap_segments = []
                overlap_length = 0
                for seg in reversed(current_chunk):
                    if overlap_length + len(seg) <= chunk_overlap:
                        overlap_segments.insert(0, seg)
                        overlap_length += len(seg)
                    else:
                        break
                
                current_chunk = overlap_segments
                current_length = overlap_length
            
            current_chunk.append(segment)
            current_length += segment_length
        
        # 添加最后一个块
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            chunks.append(Chunk(
                id=chunk_id,
                content=chunk_text,
                metadata={
                    'chunk_method': 'by_punctuation',
                    'chunk_size': chunk_size,
                    'chunk_overlap': chunk_overlap,
                    'segment_count': len(current_chunk)
                }
            ))
        
        return chunks
    
    # ==================== 智能递归切分 ====================
    
    def split_recursive(self, 
                       text: str,
                       chunk_size: int = 1000,
                       chunk_overlap: int = 200,
                       separators: Optional[List[str]] = None) -> List[Chunk]:
        """
        智能递归切分文本
        按照分隔符优先级递归切分，直到达到目标块大小
        
        Args:
            text: 输入文本
            chunk_size: 目标每块字符数
            chunk_overlap: 重叠字符数
            separators: 分隔符列表，按优先级排序（默认：双换行、单换行、句号、逗号等）
        
        Returns:
            分块列表
        """
        if not text.strip():
            return []
        
        if separators is None:
            # 默认分隔符优先级：段落 > 句子 > 短语
            separators = [
                '\n\n',      # 双换行（段落分隔）
                '\n',        # 单换行
                '。',        # 中文句号
                '.',         # 英文句号
                '！',        # 中文感叹号
                '!',         # 英文感叹号
                '？',        # 中文问号
                '?',         # 英文问号
                '；',        # 中文分号
                ';',         # 英文分号
                '，',        # 中文逗号
                ',',         # 英文逗号
            ]
        
        def _recursive_split(text: str, current_separator_index: int = 0) -> List[str]:
            """递归分割函数"""
            if len(text) <= chunk_size:
                return [text] if text.strip() else []
            
            if current_separator_index >= len(separators):
                # 如果没有更多分隔符，按长度切分
                return self._split_by_length_simple(text, chunk_size)
            
            separator = separators[current_separator_index]
            parts = text.split(separator)
            
            # 如果分割后的部分都小于目标大小，可以直接返回
            if all(len(part) <= chunk_size for part in parts if part.strip()):
                # 合并小段落，但保持重叠
                chunks = []
                current_chunk = []
                current_length = 0
                
                for part in parts:
                    if not part.strip():
                        continue
                    
                    if current_length + len(part) > chunk_size and current_chunk:
                        chunk_text = separator.join(current_chunk)
                        chunks.append(chunk_text)
                        
                        # 保留重叠
                        overlap_parts = []
                        overlap_length = 0
                        for p in reversed(current_chunk):
                            if overlap_length + len(p) <= chunk_overlap:
                                overlap_parts.insert(0, p)
                                overlap_length += len(p)
                            else:
                                break
                        
                        current_chunk = overlap_parts
                        current_length = overlap_length
                    
                    current_chunk.append(part)
                    current_length += len(part) + len(separator)
                
                if current_chunk:
                    chunks.append(separator.join(current_chunk))
                
                return chunks
            
            # 否则，继续使用下一个分隔符递归分割
            results = []
            for part in parts:
                if not part.strip():
                    continue
                if len(part) <= chunk_size:
                    results.append(part)
                else:
                    # 递归分割
                    sub_results = _recursive_split(part, current_separator_index + 1)
                    results.extend(sub_results)
            
            return results
        
        def _split_by_length_simple(text: str, size: int) -> List[str]:
            """简单按长度切分"""
            chunks = []
            start = 0
            while start < len(text):
                end = min(start + size, len(text))
                chunk = text[start:end].strip()
                if chunk:
                    chunks.append(chunk)
                start = end
            return chunks
        
        chunk_texts = _recursive_split(text)
        
        # 转换为Chunk对象
        chunks = []
        for i, chunk_text in enumerate(chunk_texts, 1):
            if chunk_text.strip():
                chunks.append(Chunk(
                    id=i,
                    content=chunk_text.strip(),
                    metadata={
                        'chunk_method': 'recursive',
                        'chunk_size': chunk_size,
                        'chunk_overlap': chunk_overlap,
                        'separators_used': separators
                    }
                ))
        
        return chunks
    
    # ==================== 按页切分 ====================
    
    def split_by_page(self, file_path: str) -> List[Chunk]:
        """
        按页切分文档（适用于PDF和Word）
        
        Args:
            file_path: 文件路径
        
        Returns:
            分块列表
        """
        ext = os.path.splitext(file_path)[1].lower()
        chunks = []
        
        if ext == '.pdf':
            if not PDF_SUPPORT:
                raise ImportError("需要安装 pdfplumber 或 PyPDF2")
            
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        text = page.extract_text()
                        if text and text.strip():
                            chunks.append(Chunk(
                                id=page_num,
                                content=text.strip(),
                                metadata={
                                    'chunk_method': 'by_page',
                                    'page_number': page_num,
                                    'total_pages': len(pdf.pages),
                                    'file_type': 'pdf'
                                }
                            ))
            except:
                # 备用方案：PyPDF2
                try:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page_num, page in enumerate(pdf_reader.pages, 1):
                            text = page.extract_text()
                            if text and text.strip():
                                chunks.append(Chunk(
                                    id=page_num,
                                    content=text.strip(),
                                    metadata={
                                        'chunk_method': 'by_page',
                                        'page_number': page_num,
                                        'total_pages': len(pdf_reader.pages),
                                        'file_type': 'pdf'
                                    }
                                ))
                except Exception as e:
                    raise ValueError(f"按页读取PDF失败: {e}")
        
        elif ext == '.docx':
            if not DOCX_SUPPORT:
                raise ImportError("需要安装 python-docx")
            
            try:
                doc = DocxDocument(file_path)
                current_page_text = []
                page_num = 1
                
                # Word文档没有真正的"页"概念，这里按段落模拟分页
                # 可以按每N个段落或每N个字符作为一页
                paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
                
                # 简单策略：每10个段落或每3000字符作为一页
                current_length = 0
                paragraph_count = 0
                page_text = []
                
                for para in paragraphs:
                    para_length = len(para)
                    page_text.append(para)
                    current_length += para_length
                    paragraph_count += 1
                    
                    # 达到"一页"的标准
                    if paragraph_count >= 10 or current_length >= 3000:
                        if page_text:
                            chunks.append(Chunk(
                                id=page_num,
                                content='\n'.join(page_text),
                                metadata={
                                    'chunk_method': 'by_page',
                                    'page_number': page_num,
                                    'paragraph_count': paragraph_count,
                                    'char_count': current_length,
                                    'file_type': 'docx'
                                }
                            ))
                            page_num += 1
                            page_text = []
                            current_length = 0
                            paragraph_count = 0
                
                # 最后一页
                if page_text:
                    chunks.append(Chunk(
                        id=page_num,
                        content='\n'.join(page_text),
                        metadata={
                            'chunk_method': 'by_page',
                            'page_number': page_num,
                            'paragraph_count': paragraph_count,
                            'char_count': current_length,
                            'file_type': 'docx'
                        }
                    ))
                
            except Exception as e:
                raise ValueError(f"按页读取DOCX失败: {e}")
        
        else:
            # 对于其他文件类型，读取文本后按固定大小模拟分页
            text = self.get_file_text(file_path)
            page_size = 3000  # 每页约3000字符
            page_num = 1
            
            for i in range(0, len(text), page_size):
                page_text = text[i:i+page_size]
                if page_text.strip():
                    chunks.append(Chunk(
                        id=page_num,
                        content=page_text.strip(),
                        metadata={
                            'chunk_method': 'by_page',
                            'page_number': page_num,
                            'char_count': len(page_text),
                            'file_type': ext[1:]
                        }
                    ))
                    page_num += 1
        
        return chunks
    
    # ==================== 统一接口 ====================
    
    def split_file(self, 
                  file_path: str,
                  method: str = 'by_length',
                  chunk_size: int = 1000,
                  chunk_overlap: int = 200) -> List[Chunk]:
        """
        统一的分块接口
        
        Args:
            file_path: 文件路径
            method: 分块方法 ('by_length', 'by_punctuation', 'recursive', 'by_page')
            chunk_size: 块大小（对于by_page方法无效）
            chunk_overlap: 重叠大小（对于by_page方法无效）
        
        Returns:
            分块列表
        """
        if method == 'by_page':
            return self.split_by_page(file_path)
        
        # 其他方法需要先读取文本
        text = self.get_file_text(file_path)
        
        if method == 'by_length':
            return self.split_by_length(text, chunk_size, chunk_overlap)
        elif method == 'by_punctuation':
            return self.split_by_punctuation(text, chunk_size, chunk_overlap)
        elif method == 'recursive':
            return self.split_recursive(text, chunk_size, chunk_overlap)
        else:
            raise ValueError(f"不支持的分块方法: {method}")

